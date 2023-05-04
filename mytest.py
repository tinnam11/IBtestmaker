import mysql.connector
import tkinter as tk
from tkinter import ttk
import PIL
from PIL import Image,ImageTk
from docx.enum.text import WD_COLOR_INDEX
from tkinter import *
from docx import Document
from docx.shared import RGBColor
from docx2pdf import convert
from tkinter import messagebox
import sys
import logging
import subprocess
import time
import os
import codecs
import edittest
from docx.shared import Inches
class MyTest:
    #contructor of the GUI for my test page
    def __init__(self, window,test_nameid): #constructor
        self.filename="temp"
        self.style = Toplevel(window)
        self.root = self.style
        # self.root.theme_use('lumen')
        # self.root.configure('', font=('', '15'))
        # self.root.configure('TCheckbutton', font=('',"15"))
        # self.root.configure('TLabelframe.Label', font=('', '15'))
        # self.root.configure('success.Outline.TButton', font=('','15'))
        # self.root.configure('custom.TFrame', background='#EBFFF1')
        # self.root.configure('TButton', font=('', '15'))
        # self.root.configure('custom.TLabel', font=('', '15'), background='#EBFFF1', foreground='#0164B4')
        self.root.title("Test Preview")
        self.root.geometry('1400x780')
        self.test_nameid=test_nameid

        self.mainframe = ttk.Frame(self.root)
        self.mainframe.pack(fill='both', expand=1)
        self.canvas = Canvas(self.mainframe)
        self.canvas.pack(side='left', fill='both',expand=1)
        self.scrollbar = ttk.Scrollbar(self.mainframe, orient='vertical', command=self.canvas.yview)
        self.scrollbar.pack(side='right', fill='y')
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.bind('<Configure>', lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.secondframe = ttk.Frame(self.canvas)
        self.canvas.create_window((0, 0), window=self.secondframe, anchor='nw')
        #function to display or not display markscheme
        self.asp = tk.IntVar()
        self.ms = tk.IntVar()

        #test name title text
        self.testlb = ttk.Label(self.secondframe, text=self.getTestname(test_nameid),font=('', 40))
        self.testlb.pack(side='top', fill='x', padx=10, pady=30)
        #Add main boxed frame surrounding all edits
        self.mainframe = ttk.Labelframe(self.secondframe, text="All actions", padding=(20, 20, 10, 10),labelanchor='nw')
        self.mainframe.pack(anchor='w', side=TOP,padx=10, pady=30, fill=X, expand=False)
        #test preview options frame
        self.testpreviewframe = ttk.Labelframe(self.mainframe, text="preview options", padding=(20, 20, 10, 10))
        self.testpreviewframe.grid(row=0, column=0, padx=10,ipadx=40)
        self.answerspacebutton = ttk.Checkbutton(self.testpreviewframe, text='Answer space', onvalue=1, offvalue=0, variable=self.asp,command=lambda:self.showresults(self.ms.get(),self.asp.get()))
        self.answerspacebutton.grid(row=0, column=0, sticky='ew', padx=5, pady=10)
        self.markschemebutton = ttk.Checkbutton(self.testpreviewframe, text='Markscheme', onvalue=1, offvalue=0,variable=self.ms, command=lambda:self.showresults(self.ms.get(),self.asp.get()))
        self.markschemebutton.grid(row=1, column=0, sticky='ew', padx=5, pady=10)
        #edits frame
        self.addonframe = ttk.Labelframe(self.mainframe, text="edits", padding=(20, 20, 10, 10))
        self.addonframe.grid(row=0, column=1,padx=10,ipadx=40)
        self.addibquestions = ttk.Button(self.addonframe, text='add IB question', style='success.Outline.TButton')
        self.addibquestions.grid(row=0, column=0, sticky='ew', padx=5, pady=10)
        self.addownquestion = ttk.Button(self.addonframe, text='your own question', style='success.Outline.TButton')
        self.addownquestion.grid(row=1, column=0, sticky='ew', padx=5, pady=10)
        self.edittestname = ttk.Button(self.addonframe, text='edit test name and description', style='success.Outline.TButton', command=lambda:self.etest(test_nameid))
        self.edittestname.grid(row=2, column=0, sticky='ew', padx=5, pady=10)
        # adding download features
        self.downloadframe = ttk.Labelframe(self.mainframe, text="download options", padding=(20, 20, 10, 10))
        self.downloadframe.grid(row=0, column=2, padx=10, ipadx=50)
        self.downloadpdf = ttk.Button(self.downloadframe, text='download as PDF', command=self.docxtopdf)
        self.downloadpdf.grid(row=0, column=0, sticky='ew', padx=5, pady=10)
        self.downloaddocx = ttk.Button(self.downloadframe, text='download as docx', command= lambda:self.createtextdoc(self.ms.get(),self.asp.get(), test_nameid))
        self.downloaddocx.grid(row=1, column=0, sticky='ew', padx=5, pady=10)
        self.printfile = ttk.Button(self.downloadframe, text='print file')
        self.printfile.grid(row=2, column=0, sticky='ew', padx=5, pady=10)
        # questions text
        # self.temptest = []
        # self.temptest = self.getTest(test_nameid)
        #creating new list from getTest function (separates the list to 3 indivual items)
        self.list1, self.list2, self.list3,self.list4, self.list5= self.getTest(test_nameid)
        #creating a list to define the object number
        self.msframe = []
        self.mslb=[]
        self.aspframe=[]
        self.asplb=[]
        print("image")
        print(self.list5)



        #looping through the number of question in list to create new label frames + questions
        for q in range(0, len(self.list1)):
            #creating frames for each question
            self.bigframe = ttk.Labelframe(self.secondframe, padding=(20, 20, 10, 10))
            self.bigframe.pack(side=BOTTOM, fill=X, padx=10, pady=30, expand=False)
            #creating question frame inside
            self.questionframe = ttk.Frame(self.bigframe)
            self.questionframe.grid(sticky='ew', row=0, column=0)
            #creating question label
            self.qlb = ttk.Label(self.questionframe, text=('Question:' + '\n' + self.list1[q]), style='TLabel',
                                     font=('', 15))
            self.qlb.grid(sticky='ew', row=0, column=0)
            self.im=self.list5[q]
            if len(self.im)>0:
                self.image1 = PIL.Image.open(self.im)
                self.resizeim=self.image1.resize((400,200))
                self.test = ImageTk.PhotoImage(self.resizeim)
            #
            # # self.imgToInsert = ImageTk.PhotoImage(file=self.im)
            # # self.pilimg=Image.open(self.im)
            # # self.img=ImageTk.PhotoImage(self.pilimg)
                self.image=ttk.Label(self.questionframe, image=self.test)
                self.image.grid(sticky='ew', row=1, column=0)
            else:
                pass
            #creating answer frame inside the bigframe
            self.answerframe = ttk.Frame(self.bigframe)
            self.answerframe.grid(sticky='ew', row=1, column=0)
            #answer label
            self.aclb = ttk.Label(self.answerframe,
                                      text=(self.list2[q]),
                                      style='TLabel', font=('', 15))
            self.aclb.grid(sticky='ew', row=0, column=0)
            #creating the markcheme frame + label by appending it into a list to make the obejct unique
            self.msframe.append(ttk.Frame(self.bigframe, style='custom.TFrame'))
            self.mslb.append(ttk.Label(self.msframe[q], text=('Markscheme:' + '\n' + self.list3[q]),
                                  font=('', 15), style='custom.TLabel'))
            # creating the answer space frame + label by appending it into a list to make the obejct unique
            self.aspframe.append(ttk.Frame(self.bigframe))
            self.asplb.append(ttk.Label(self.aspframe[q], text=self.list4[q],
                                   font=('', 15), style='TLabel'))
        self.root.mainloop()
    #creating a docx file from the users selected preview
    def createtextdoc(self,markscheme, answerspace, test_nameid):
        # creating a file
        document = Document()
        #write into a document with markscheme and answer space
        if markscheme==1 and answerspace ==1:
            #name the document stating that MS is included
            self.filename="MS "+self.getTestname(test_nameid)
            # loop in list to add each question + markscheme + answer space in order
            for q in range(0, len(self.list1)):
                #write in a new line for each elements
                document.add_paragraph(str(q+1)+'.' + '\n' + self.list1[q] + '\n' + '\n')
                self.im = self.list5[q]
                if len(self.im) > 0:
                    document.add_picture(self.im[q])
                else:
                    pass
                document.add_paragraph(self.list2[q] + '\n' + '\n')
                mark = document.add_paragraph().add_run('Markscheme:                                                                                                                                                         ' + '\n' + self.list3[q]+'                                                                                                                                                                                     ' + '\n')
                document.add_paragraph('\n'+self.list4[q]+'\n')
                # highlight the markscheme and change font color
                font = mark.font
                font.color.rgb = RGBColor(0,99,180)
                font.highlight_color = WD_COLOR_INDEX.GRAY_25
        elif markscheme==1 and answerspace==0:
            self.filename = "MS " + self.getTestname(test_nameid)
            self.f = open(self.filename + ".txt", "w+")
            #loop in list to add each question + markscheme in order
            for q in range(0, len(self.list1)):
                document.add_paragraph(str(q + 1) + '.' + '\n' + self.list1[q] + '\n' + '\n')
                self.im = self.list5[q]
                if len(self.im) > 0:
                    document.add_picture(self.im)
                else:
                    pass
                document.add_paragraph(self.list2[q] + '\n' + '\n')
                mark = document.add_paragraph().add_run('Markscheme:' + '\n' + self.list3[q] + '\n')
                font = mark.font
                #highlight the markscheme and change font color
                font.highlight_color = WD_COLOR_INDEX.GRAY_25
                font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        #creating file with answer space
        elif markscheme==0 and answerspace==1:
            self.filename = self.getTestname(test_nameid)
            self.f = open(self.filename + ".txt", "w+")
            #loop through the list to write question and answer space in order
            for q in range(0, len(self.list1)):
                document.add_paragraph(str(q + 1) + '.' + '\n' + self.list1[q] + '\n' + '\n')
                self.im = self.list5[q]
                if len(self.im) > 0:
                    document.add_picture(self.im)
                else:
                    pass
                document.add_paragraph(self.list2[q] + '\n' + '\n')
                document.add_paragraph('\n' + self.list4[q] + '\n')
        else:
            self.filename = self.getTestname(test_nameid)
            self.f = open(self.filename + ".txt", "w+")
            for q in range(0, len(self.list1)):
                document.add_paragraph(str(q + 1) + '.' + '\n' + self.list1[q] + '\n' + '\n')
                self.im = self.list5[q]
                if len(self.im) > 0:
                    document.add_picture(self.im)
                else:
                    pass
                document.add_paragraph(self.list2[q] + '\n' + '\n')

        document.save(self.filename +'.docx')
        convert(self.filename+'.docx')
    # get operating system of the laptop
    def get_platform(self):
        if sys.platform == 'linux':
            try:
                proc_version = open('/proc/version').read()
                if 'Microsoft' in proc_version:
                    return 'wsl'
            except:
                pass
        return sys.platform
    # open default app with the chosen file to display the test file
    def open_with_default_app(self,filename):
        platform = self.get_platform()
        if platform == 'darwin':
            subprocess.call(('open', filename))
        elif platform in ['win64', 'win32']:
            os.startfile(filename.replace('/', '\\'))
        elif platform == 'wsl':
            subprocess.call('cmd.exe /C start'.split() + [filename])
        else:  # linux variants
            subprocess.call(('xdg-open', filename))
    #converts the docx to a pdf file
    def docxtopdf(self):
        if self.filename != "temp":
            #opens the docx file in a default app
            self.open_with_default_app('/Users/tinnam/Downloads/IAproject/'+self.filename+'.docx')
            time.sleep(5)
            #converts the file to pdf
            convert('/Users/tinnam/Downloads/IAproject/'+self.filename+'.docx')

        else:
            self.message=messagebox.showwarning("error", "a document has to be generated first, press 'download as docx'")

        #
    #function to allow the checkbutton in preview options to work
    def showresults(self, markscheme, answerspace):
        for q in range(0, len(self.list1)):
            #if marscheme is presses, show markscheme by griddinng pre existing objects in the list
            if markscheme == 1:
                self.msframe[q].grid(sticky='ew', row=3, column=0, columnspan=10, ipadx=200)
                self.mslb[q].grid(row=0, column=0, sticky='ew', columnspan=10)
            #if not checked or unchecked, un grid the object
            elif markscheme==0:
                self.msframe[q].grid_forget()
            if answerspace == 1:
                self.aspframe[q].grid(sticky='ew', row=2, column=0, columnspan=10)
                self.asplb[q].grid(row=0, column=0, sticky='ew', columnspan=10,pady=10)
            elif answerspace==0:
                self.aspframe[q].grid_forget()
    #getting the testname from the selected test in homepage treeview option
    def getTestname(self, testid):
        self.mydb = mysql.connector.connect(
        host="localhost",
        user="root	",
        password="",
        # port=3307,
        database="test"
    )
        self.mycursor = self.mydb.cursor()
        self.sql = "select testname from mytests where testid = '" + testid + "'"
        print(self.sql)
        self.mycursor.execute(self.sql)
        self.results = self.mycursor.fetchone()
        print(self.results[0])
        return (self.results[0])
    #getting the question in the test from the selected test in homepage and appends the q,ms,asp into diff lists
    def getTest(self, testid):
        self.mydb = mysql.connector.connect(
        host="localhost",
        user="root	",
        password="",
        # port=3307,
        database="test"
    )
        self.mycursor = self.mydb.cursor()
        self.sql="select list_questionid from mytests where testid = '"+testid+"'"
        print(self.sql)
        self.mycursor.execute(self.sql)
        self.results = self.mycursor.fetchall() #Output -> row
        self.temp = self.results[0][0]
        self.tempstr=""
        if "," in self.temp:
            self.listOfquestion=self.temp.split(',')
            print(self.listOfquestion)
            for i in range (len(self.listOfquestion)):
                if i == len(self.listOfquestion)-1:
                    self.tempstr = self.tempstr + "'" + self.listOfquestion[i]+ "'"
                else:
                    self.tempstr = self.tempstr + "'" + self.listOfquestion[i]+ "',"
            self.temp = self.tempstr
        else:
            self.temp = "'" + self.temp + "'"
        self.sql="select * from newpastpaper where No in ("+self.temp+")"
        print(self.sql)
        self.mycursor.execute(self.sql)
        self.results = self.mycursor.fetchall()
        self.qlist=[]
        self.aclist=[]
        self.mslist=[]
        self.aslist=[]
        self.ilist=[]
        #appending question text
        for i in self.results:
            self.qlist.append(i[1])
        # appending answer choice text
        for i in self.results:
            self.aclist.append(i[2])
        # appending markscheme text
        for i in self.results:
            self.mslist.append(i[3])
        #appending answer space
        for i in self.results:
            self.aslist.append(i[15])
        # print(self.mslist)
        # print(self.qlist)
        for i in self.results:
            self.ilist.append(i[4])
        return self.qlist, self.aclist, self.mslist,self.aslist, self.ilist
    def etest(self,testid):
        self.A = edittest.Edittest(self.root, testid)




