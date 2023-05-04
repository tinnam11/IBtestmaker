import os

import docx2pdf
# from docx import Document
# #
# doc = Document()
# with open("pdftest.txt", 'r', encoding='utf-8') as file:
#     doc.add_paragraph(file.read())
# doc.save("covertdoc.docx")

from docx2pdf import convert
import os
from docx import Document

document = Document('/Users/tinnam/Downloads/IAproject/covertdoc.docx')
document.save('/Users/tinnam/Downloads/IAproject/hahaha.docx')
# os.startfile('/Users/tinnam/Downloads/IAproject/covertdoc.docx')
# os.sta

#convert a single docx file to pdf file in same directory
# convert('covertdoc.docx')
# print('complete')
#
# #convert docx to pdf specifying input & output paths
# document=Document('/Users/tinnam/Downloads/IAproject/covertdoc.docx')
# doc=document.open('/Users/tinnam/Downloads/IAproject/covertdoc.docx')
# document.saveas('/Users/tinnam/Downloads/IAproject/covertdoc.pdf')
# document.

import os
import sys


# os.startfile('start /Users/tinnam/Downloads/IAproject/covertdoc.docx')
# convert(r'/Users/tinnam/Downloads/IAproject/MS atomic structure test.docx')
# convert('/Users/tinnam/Downloads/IAproject/' + self.filename + '.docx',
#         '/Users/tinnam/Downloads/IAproject/' + self.filename + '.pdf')
# convert('/Users/tinnam/Downloads/IAproject')

# from docx import Document
# from docx.shared import RGBColor
# document = Document()
# run = document.add_paragraph().add_run('some text')
# font = run.font
# font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
# p=document.add_paragraph('aaa')
# document.save('demo1.docx')
#
#
# import win32com.client as client
#
#
# def convert_to_pdf(filepath:str):
#     """Save a pdf of a docx file."""
#     try:
#         word = client.DispatchEx("Word.Application")
#         target_path = filepath.replace(".docx", r".pdf")
#         word_doc = word.Documents.Open(filepath)
#         word_doc.SaveAs(target_path, FileFormat=17)
#         word_doc.Close()
#     except Exception as e:
#             raise e
#     finally:
#             word.Quit()